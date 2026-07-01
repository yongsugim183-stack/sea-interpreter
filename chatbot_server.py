"""
Intel AI for Workforce 교육과정 안내 챗봇 서버
개인비서 챗봇 + Google Calendar 연동 + 커리큘럼(PDF/Word) 분석
동시통역 서버(server.py)와 완전히 분리된 독립 프로세스로 실행됨
"""

import asyncio
import json
import os
import re as _re
from datetime import datetime as _dt, timedelta as _td, timezone as _tz
from pathlib import Path as _Path

from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel

app = FastAPI(title="개인비서 챗봇 API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.get("/health")
async def health():
    return {"status": "ok"}


@app.get("/chatbot")
async def chatbot_page():
    return FileResponse("chatbot.html")


# ── 커리큘럼(PDF/Word) 업로드 및 분석 ─────────────────────────────────────

CURRICULUM_FILE = "curriculum_data.json"


def _load_curriculum() -> dict:
    if os.path.exists(CURRICULUM_FILE):
        with open(CURRICULUM_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}


def _save_curriculum(data: dict):
    with open(CURRICULUM_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


# 한국어 조사/어미 제거 패턴 (긴 것부터 우선 매칭)
_KO_JOSA = _re.compile(
    r'(이라고|이라는|에서는|으로는|으로서|부터는|까지는|에게는|로부터|'
    r'에서|으로|이라|이고|이며|이나|이나|라고|라는|라도|에게|한테|보다|'
    r'에도|의해|처럼|만큼|까지|부터|조차|마저|이든|이면|이야|이어|'
    r'에서|에는|에도|에만|에게|의|을|를|이|가|은|는|도|로|으|와|과|'
    r'아|야|이야|랑|이랑|나|이나|며|이며|고|이고|서|에서)$'
)


def _normalize_ko(word: str) -> str:
    """한국어 단어에서 조사/어미를 제거해 어근 추출"""
    if len(word) <= 2:
        return word
    m = _KO_JOSA.search(word)
    if m:
        stem = word[:m.start()]
        if len(stem) >= 2:
            return stem
    return word


def _tokenize(text: str) -> list[str]:
    """텍스트를 토큰 목록으로 분리 (한국어 조사 제거 포함)"""
    raw = _re.findall(r'[가-힣]{2,}|[a-zA-Z0-9®]{2,}', text)
    result = []
    for w in raw:
        norm = _normalize_ko(w)
        if len(norm) >= 2:
            result.append(norm)
        # 원본도 추가 (조사 포함 원형으로도 검색되도록)
        if norm != w and len(w) >= 2:
            result.append(w)
    return result


def _build_keywords(text: str) -> list[str]:
    tokens = _tokenize(text)
    freq: dict[str, int] = {}
    for t in tokens:
        freq[t] = freq.get(t, 0) + 1
    return [w for w, _ in sorted(freq.items(), key=lambda x: -x[1])[:30]]


def _analyze_docx(docx_bytes: bytes) -> dict:
    """Word(.docx) 파일 분석 - 스타일 기반 제목 감지로 PDF보다 정확"""
    import docx as _docx
    import io

    doc = _docx.Document(io.BytesIO(docx_bytes))

    HEADING_STYLES = {"heading 1", "heading 2", "heading 3",
                      "제목 1", "제목 2", "제목 3",
                      "머리글 1", "머리글 2", "머리글 3"}

    def is_heading(para) -> bool:
        style_name = (para.style.name or "").lower()
        if style_name in HEADING_STYLES:
            return True
        # Bold + 짧은 줄도 제목으로
        text = para.text.strip()
        if not text or len(text) > 60:
            return False
        if all(run.bold for run in para.runs if run.text.strip()):
            return True
        return False

    sections: list[dict] = []
    cur_heading = "일반 내용"
    cur_body: list[str] = []

    def flush_section(heading, body_lines):
        body = "\n".join(body_lines).strip()
        if len(body) >= 10:
            sections.append({
                "heading": heading,
                "body": body,
                "page": 1,
                "keywords": _build_keywords(heading + " " + body),
            })

    all_lines: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        all_lines.append(text)
        if is_heading(para):
            flush_section(cur_heading, cur_body)
            cur_heading = text
            cur_body = []
        else:
            cur_body.append(text)

    # 표(table) 내용도 추가
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                cur_body.append(row_text)
                all_lines.append(row_text)

    flush_section(cur_heading, cur_body)

    # 섹션이 너무 적으면 단락 단위 재분할
    full_text = "\n".join(all_lines)
    if len(sections) < 3:
        sections = []
        for para in _re.split(r'\n{2,}', full_text):
            para = para.strip()
            if len(para) < 15:
                continue
            first = para.splitlines()[0].strip()
            heading = first if len(first) <= 40 else "내용"
            sections.append({
                "heading": heading,
                "body": para,
                "page": 1,
                "keywords": _build_keywords(para),
            })

    # 슬라이딩 윈도우 청크
    CHUNK_LINES = 15
    STEP = 7
    chunks: list[dict] = []
    for i in range(0, max(1, len(all_lines) - CHUNK_LINES + 1), STEP):
        chunk_text = "\n".join(all_lines[i:i + CHUNK_LINES])
        chunks.append({"body": chunk_text, "keywords": _build_keywords(chunk_text)})
    if all_lines and (not chunks or all_lines[-1] not in chunks[-1]["body"]):
        tail = "\n".join(all_lines[-(CHUNK_LINES // 2):])
        if tail.strip():
            chunks.append({"body": tail, "keywords": _build_keywords(tail)})

    return {
        "full_text": full_text,
        "page_count": len(doc.sections),   # Word의 섹션 수 (페이지 수 대용)
        "sections": sections,
        "section_count": len(sections),
        "chunks": chunks,
    }


def _analyze_pdf(pdf_bytes: bytes) -> dict:
    import fitz  # PyMuPDF

    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    page_count = len(doc)

    # ── 1단계: 스팬(span) 단위 텍스트 추출 (폰트 크기·굵기 포함) ──────────
    # 각 스팬: {text, size, bold, bbox, page}
    spans: list[dict] = []
    for page_num, page in enumerate(doc):
        blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
        for block in blocks:
            if block.get("type") != 0:   # 0 = 텍스트 블록
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    text = span["text"].strip()
                    if not text:
                        continue
                    flags = span.get("flags", 0)
                    spans.append({
                        "text": text,
                        "size": round(span["size"], 1),
                        "bold": bool(flags & 2**4),   # bit 4 = bold
                        "bbox": span["bbox"],
                        "page": page_num + 1,
                    })

    doc.close()

    if not spans:
        return {"full_text": "", "page_count": page_count,
                "sections": [], "section_count": 0, "chunks": []}

    # ── 2단계: 폰트 크기 분포로 제목 크기 임계값 결정 ────────────────────
    sizes = [s["size"] for s in spans]
    median_size = sorted(sizes)[len(sizes) // 2]
    # 중앙값보다 1.1배 이상 크거나 Bold 이면 제목 후보
    HEADING_RATIO = 1.1

    # ── 3단계: 스팬을 논리 줄(line)로 병합 ──────────────────────────────
    # 같은 페이지 + Y좌표가 거의 같으면 한 줄로 합산
    lines_merged: list[dict] = []
    for sp in spans:
        if lines_merged:
            prev = lines_merged[-1]
            same_line = (
                prev["page"] == sp["page"]
                and abs(prev["bbox"][1] - sp["bbox"][1]) < 3
            )
            if same_line:
                prev["text"] += " " + sp["text"]
                prev["size"] = max(prev["size"], sp["size"])
                prev["bold"] = prev["bold"] or sp["bold"]
                prev["bbox"] = (
                    min(prev["bbox"][0], sp["bbox"][0]),
                    min(prev["bbox"][1], sp["bbox"][1]),
                    max(prev["bbox"][2], sp["bbox"][2]),
                    max(prev["bbox"][3], sp["bbox"][3]),
                )
                continue
        lines_merged.append({
            "text": sp["text"],
            "size": sp["size"],
            "bold": sp["bold"],
            "bbox": sp["bbox"],
            "page": sp["page"],
        })

    # ── 4단계: 제목 vs 본문 판별 ─────────────────────────────────────────
    def is_heading(line: dict) -> bool:
        t = line["text"].strip()
        if not t or len(t) > 60:
            return False
        # 폰트 크기 또는 Bold
        if line["size"] >= median_size * HEADING_RATIO or line["bold"]:
            return True
        # 짧고 의미 있는 한국어 줄
        if len(t) <= 20 and _re.match(r'^[가-힣®\s\d\.\)\-·×,()]+$', t):
            return True
        return False

    # ── 5단계: 섹션 구조화 ───────────────────────────────────────────────
    sections: list[dict] = []
    cur_heading = "일반 내용"
    cur_body: list[str] = []
    cur_page = 1

    def flush_section(heading, body_lines, page):
        body = "\n".join(body_lines).strip()
        if len(body) >= 10:
            sections.append({
                "heading": heading,
                "body": body,
                "page": page,
                "keywords": _build_keywords(heading + " " + body),
            })

    for ln in lines_merged:
        t = ln["text"].strip()
        if not t:
            continue
        if is_heading(ln):
            flush_section(cur_heading, cur_body, cur_page)
            cur_heading = t
            cur_body = []
            cur_page = ln["page"]
        else:
            cur_body.append(t)
    flush_section(cur_heading, cur_body, cur_page)

    # 섹션이 너무 적으면 단락 단위로 재분할
    full_text = "\n".join(ln["text"] for ln in lines_merged)
    if len(sections) < 3:
        sections = []
        for para in _re.split(r'\n{2,}', full_text):
            para = para.strip()
            if len(para) < 15:
                continue
            first = para.splitlines()[0].strip()
            heading = first if len(first) <= 40 else "내용"
            sections.append({
                "heading": heading,
                "body": para,
                "page": 1,
                "keywords": _build_keywords(para),
            })

    # ── 6단계: 슬라이딩 윈도우 청크 (섹션 경계와 무관한 연속 문맥 보존) ──
    all_lines = [ln["text"].strip() for ln in lines_merged if ln["text"].strip()]
    CHUNK_LINES = 15
    STEP = 7
    chunks: list[dict] = []
    for i in range(0, max(1, len(all_lines) - CHUNK_LINES + 1), STEP):
        chunk_text = "\n".join(all_lines[i:i + CHUNK_LINES])
        chunks.append({"body": chunk_text, "keywords": _build_keywords(chunk_text)})
    # 마지막 잔여 줄
    if all_lines and (not chunks or all_lines[-1] not in chunks[-1]["body"]):
        tail = "\n".join(all_lines[-(CHUNK_LINES // 2):])
        if tail.strip():
            chunks.append({"body": tail, "keywords": _build_keywords(tail)})

    return {
        "full_text": full_text,
        "page_count": page_count,
        "sections": sections,
        "section_count": len(sections),
        "chunks": chunks,
    }


@app.post("/api/upload-curriculum")
async def upload_curriculum(file: UploadFile = File(...)):
    fname = file.filename.lower()
    if not (fname.endswith(".pdf") or fname.endswith(".docx")):
        raise HTTPException(status_code=400, detail="PDF 또는 Word(.docx) 파일만 업로드 가능합니다.")
    file_bytes = await file.read()
    if len(file_bytes) > 20 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="파일 크기는 20MB 이하여야 합니다.")
    try:
        loop = asyncio.get_event_loop()
        if fname.endswith(".docx"):
            data = await loop.run_in_executor(None, _analyze_docx, file_bytes)
        else:
            data = await loop.run_in_executor(None, _analyze_pdf, file_bytes)
        data["filename"] = file.filename
        _save_curriculum(data)
        return {
            "success": True,
            "filename": file.filename,
            "page_count": data["page_count"],
            "section_count": data["section_count"],
            "text_length": len(data["full_text"]),
            "sections": data["sections"],
            "chunks": data["chunks"],
            "full_text": data["full_text"],
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"파일 분석 오류: {str(e)}")


@app.get("/api/curriculum")
async def get_curriculum():
    data = _load_curriculum()
    if not data:
        return {"loaded": False}
    return {
        "loaded": True,
        "filename": data.get("filename", ""),
        "page_count": data.get("page_count", 0),
        "section_count": data.get("section_count", 0),
        "text_length": len(data.get("full_text", "")),
        "sections": data.get("sections", []),
        "chunks": data.get("chunks", []),
        "full_text": data.get("full_text", ""),
    }


# ── Google Calendar 연동 ─────────────────────────────────────────────────

CALENDAR_TOKEN_FILE = "calendar_token.json"
CALENDAR_CREDS_FILE = "credentials.json"
CALENDAR_DATA_FILE = "calendar_data.json"

GCAL_SCOPES = ["https://www.googleapis.com/auth/calendar.readonly"]


def _ensure_credentials_file():
    """credentials.json이 없고 GOOGLE_OAUTH_CREDENTIALS_JSON 환경변수가 있으면 파일로 기록.
    (Render 등 원격 배포 시 저장소에 시크릿 파일을 커밋하지 않기 위함)"""
    if _Path(CALENDAR_CREDS_FILE).exists():
        return
    raw = os.environ.get("GOOGLE_OAUTH_CREDENTIALS_JSON")
    if raw:
        _Path(CALENDAR_CREDS_FILE).write_text(raw, encoding="utf-8")


def _load_calendar_data() -> dict:
    p = _Path(CALENDAR_DATA_FILE)
    if p.exists():
        try:
            return json.loads(p.read_text(encoding="utf-8"))
        except Exception:
            pass
    return {}


def _save_calendar_data(data: dict):
    _Path(CALENDAR_DATA_FILE).write_text(
        json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8"
    )


def _get_credentials():
    """저장된 토큰으로 Credentials 반환. 없으면 None."""
    from google.oauth2.credentials import Credentials
    from google.auth.transport.requests import Request
    p = _Path(CALENDAR_TOKEN_FILE)
    if not p.exists():
        return None
    creds = Credentials.from_authorized_user_file(str(p), GCAL_SCOPES)
    if creds and creds.expired and creds.refresh_token:
        try:
            creds.refresh(Request())
            p.write_text(creds.to_json(), encoding="utf-8")
        except Exception:
            return None
    return creds if creds and creds.valid else None


def _fetch_events(creds) -> list[dict]:
    """Google Calendar에서 향후 90일 + 과거 30일 이벤트 조회"""
    from googleapiclient.discovery import build
    service = build("calendar", "v3", credentials=creds)
    now = _dt.now(_tz.utc)
    time_min = (now - _td(days=30)).isoformat()
    time_max = (now + _td(days=90)).isoformat()
    result = service.events().list(
        calendarId="primary",
        timeMin=time_min,
        timeMax=time_max,
        maxResults=250,
        singleEvents=True,
        orderBy="startTime",
    ).execute()
    events = result.get("items", [])
    parsed = []
    for e in events:
        start = e.get("start", {})
        end = e.get("end", {})
        start_str = start.get("dateTime", start.get("date", ""))
        end_str = end.get("dateTime", end.get("date", ""))
        parsed.append({
            "id": e.get("id", ""),
            "title": e.get("summary", "(제목 없음)"),
            "start": start_str,
            "end": end_str,
            "location": e.get("location", ""),
            "description": e.get("description", ""),
            "allDay": "date" in start,
        })
    return parsed


def _events_to_knowledge(events: list[dict]) -> dict:
    """이벤트 목록을 챗봇 지식 베이스(sections + chunks)로 변환"""
    sections = []
    all_lines = []

    for e in events:
        start = e["start"]
        try:
            if "T" in start:
                dt_obj = _dt.fromisoformat(start.replace("Z", "+00:00"))
                kst = dt_obj.astimezone(_tz(offset=_td(hours=9)))
                date_str = kst.strftime("%Y년 %m월 %d일 %H:%M")
            else:
                dt_obj = _dt.fromisoformat(start)
                date_str = dt_obj.strftime("%Y년 %m월 %d일 (종일)")
        except Exception:
            date_str = start

        lines = [f"일정: {e['title']}", f"날짜/시간: {date_str}"]
        if e.get("location"):
            lines.append(f"장소: {e['location']}")
        if e.get("description"):
            lines.append(f"내용: {e['description'][:300]}")
        body = "\n".join(lines)
        all_lines.extend(lines)
        all_lines.append("")

        sections.append({
            "heading": e["title"],
            "body": body,
            "page": 1,
            "keywords": _build_keywords(body),
        })

    full_text = "\n".join(all_lines)

    # 슬라이딩 윈도우 청크
    flat = [l for l in all_lines if l.strip()]
    CHUNK_LINES, STEP = 15, 7
    chunks = []
    for i in range(0, max(1, len(flat) - CHUNK_LINES + 1), STEP):
        chunk_text = "\n".join(flat[i:i + CHUNK_LINES])
        chunks.append({"body": chunk_text, "keywords": _build_keywords(chunk_text)})
    if flat and (not chunks or flat[-1] not in chunks[-1]["body"]):
        tail = "\n".join(flat[-(CHUNK_LINES // 2):])
        if tail.strip():
            chunks.append({"body": tail, "keywords": _build_keywords(tail)})

    return {
        "source": "google_calendar",
        "synced_at": _dt.now().isoformat(),
        "event_count": len(events),
        "full_text": full_text,
        "page_count": 1,
        "sections": sections,
        "section_count": len(sections),
        "chunks": chunks,
    }


@app.get("/api/calendar/status")
async def calendar_status():
    creds = _get_credentials()
    data = _load_calendar_data()
    return {
        "connected": creds is not None,
        "has_data": bool(data),
        "event_count": data.get("event_count", 0),
        "synced_at": data.get("synced_at", ""),
    }


_auth_in_progress = False


def _run_installed_auth() -> bool:
    """InstalledAppFlow로 브라우저 인증 후 토큰 저장 (별도 스레드에서 실행)"""
    from google_auth_oauthlib.flow import InstalledAppFlow
    flow = InstalledAppFlow.from_client_secrets_file(CALENDAR_CREDS_FILE, scopes=GCAL_SCOPES)
    creds = flow.run_local_server(port=0, open_browser=True)
    _Path(CALENDAR_TOKEN_FILE).write_text(creds.to_json(), encoding="utf-8")
    events = _fetch_events(creds)
    data = _events_to_knowledge(events)
    _save_calendar_data(data)
    return True


@app.get("/api/calendar/auth")
async def calendar_auth():
    """데스크톱 앱 OAuth 흐름 시작 — 브라우저 열림, 완료 시 토큰 저장"""
    global _auth_in_progress
    _ensure_credentials_file()
    if not _Path(CALENDAR_CREDS_FILE).exists():
        raise HTTPException(status_code=400, detail="credentials.json 파일이 없습니다.")
    if _auth_in_progress:
        raise HTTPException(status_code=409, detail="이미 인증 진행 중입니다.")
    _auth_in_progress = True
    try:
        loop = asyncio.get_event_loop()
        await loop.run_in_executor(None, _run_installed_auth)
        return {"success": True, "message": "인증 완료"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"인증 오류: {str(e)}")
    finally:
        _auth_in_progress = False


@app.get("/api/calendar/sync")
async def calendar_sync():
    """이벤트 재동기화"""
    creds = _get_credentials()
    if not creds:
        raise HTTPException(status_code=401, detail="Google Calendar 연결이 필요합니다.")
    loop = asyncio.get_event_loop()
    events = await loop.run_in_executor(None, _fetch_events, creds)
    data = _events_to_knowledge(events)
    _save_calendar_data(data)
    return {
        "success": True,
        "event_count": len(events),
        "synced_at": data["synced_at"],
        "sections": data["sections"],
        "chunks": data["chunks"],
        "full_text": data["full_text"],
    }


@app.get("/api/calendar/data")
async def calendar_data():
    data = _load_calendar_data()
    if not data:
        return {"loaded": False}
    return {
        "loaded": True,
        "event_count": data.get("event_count", 0),
        "synced_at": data.get("synced_at", ""),
        "section_count": data.get("section_count", 0),
        "sections": data.get("sections", []),
        "chunks": data.get("chunks", []),
        "full_text": data.get("full_text", ""),
    }


# ── 교육과정 오리엔테이션 챗봇 ────────────────────────────────────────────

CURRICULUM_SYSTEM_PROMPT = """당신은 한국생산성본부(KPC)의 'Intel® AI for Workforce - powered by KPC' 과정 안내 챗봇입니다.
수강 희망자들이 교육과정, 일정, 신청 방법, 특전 등에 대해 궁금한 것을 친절하고 정확하게 안내해주세요.

[과정 기본 정보]
- 과정명: Intel® AI for Workforce - powered by KPC (미래인재 육성을 위한 AI 강사양성과정)
- 주관: 한국생산성본부(KPC) × 인텔(Intel) × Mind Canvas 공동 운영
- 1차 교육 기간: 2025년 11월 16일(일) ~ 11월 23일(일), 총 6일 27시간
- 교육 장소: 한국생산성본부 강의장 (서울 종로구 새문안로5가길 32)
- 교육비: 250만원 (교재·중식 포함)
- 담당자 문의: 02-3489-475

[신청 자격]
- AI 강사 또는 IT 강사 경력 3년 이상
- 개발 경력 + 강의 경력 합산 3년 이상
- 비전공자 중 프로그래밍 가능자

[신청 기간 및 방법]
- 신청 기간: 2025년 10월 29일(수) ~ 11월 12일(수)
- 방법 1 (홈페이지): 본부 홈페이지 회원가입 → 과정 검색 → 신청하기 → 결제
- 방법 2 (구글폼): QR 코드 접속 후 폼 작성 (강사지원신청 바로가기)
- 신청 후 경력·학력 등 서류 제출, 심사 후 개별 통보

[수강 특전]
1. 인텔 정식 강사 인증서 발급
2. Intel® AI Future Workforce 과정 Full Contents 사용권
   - AI 미래인력 양성 5단계 숙련 과정 (42개 모듈, 248H)
   - 인식제고(6H) → 기초(57H) → 경험(161H) → 성과(24H) → 실무
3. 인텔 국내 오픈과정(K-Digital Training, 대학 프로그램 등) 강사 추천
4. 강의 경력에 따른 추가 강의 배정, 수료생에게 인텔 인증서 제공

[교육 일정 상세]
Day 1 (11.16 일, 10:00~17:00)
- 강사양성교육 목표 및 Q&A
- 성공적인 교수법 (1), (2) – 피칭 연습, 멘토링, 프리젠테이션 스킬
- Intel AI for Future Workforce 소개 (전체 개요 및 학습 목표)
- Awareness(m1~m6): AI 개념·활용·포용성·미래·영향·목표 설정
- AI Project Cycle(m7), Python(m9~m10): AI 프로젝트 6단계, 파이썬 기초
- Visualization(m12): 데이터 시각화 및 대시보드 작성 실습

Day 2 (11.17 월, 19:00~22:00)
- Common Trade App-NoCoding(m8): 노코드 도구로 산업별 AI 사례 체험
- Common Trade App-Coding(m11): Python으로 실제 산업형 AI 모델 구현
- Career Growth Skills(m13~m15): 인텔 디자인씽킹·시스템씽킹·기업가정신을 통한 문제해결

Day 3 (11.18 화, 19:00~22:00)
- ML/DI Techniques(m16): 지도·비지도·강화학습, 딥러닝 개념과 실습
- Statistical Data(m17~m19): 데이터 처리·통계 모델링 및 대출 승인 예제
- Computer Vision(m20~m22): OpenCV, OpenVINO 기반 영상처리·이미지 분류·폐렴 탐지

Day 4 (11.19 수, 19:00~22:00)
- Social Emotional Skills(m28), AI Ethics(m29): 감정지능·공감기술·AI 윤리 사례
- AI Project Pitfall(m30), 과제 제안: AI 프로젝트 실패 요인 및 윤리적 고려
- Natural Language Process(m24~m27): 텍스트 전처리·분류·챗봇 개발 및 호텔 챗봇 실습

Day 5 (11.20 목, 19:00~22:00)
- Introduction to Generative AI(m31): 생성형 AI 원리와 텍스트·이미지·음성 생성
- Basics of Text, Image, Video, Audio Generation(m32): 통합 활용 및 생산성 향상 실습
- Wrap up & 과제 제안: 학습 내용 정리 및 프로젝트 주제 제안

Day 6 (11.23 일, 10:00~17:00)
- Prompt Engineering(m33): 효과적인 프롬프트 설계 원리와 실습
- Gen AI Project Lifecycle & Tools(m34): 생성형 AI 개발 사이클과 주요 도구 활용
- Introduction to AI Agents(m35): AI 에이전트 구조·유형·워크플로우 설계 및 실습
- Vibe Coding Tools (Google AI Studio, Firebase Studio): 실습형 프로젝트 진행
- Project 수업 진행(Demonstration)
- 커리큘럼 운영방안 소규모 발표 및 공유
- 과정형 평가: 수업 교육역량 평가 실시, 동영상 평가

[오시는 길]
- 지하철: 3호선 경복궁역 6번 출구 아케이드 통과 지하 1층 / 5호선 광화문역 1번 출구 150m
- 버스: 광화문, 적선동, 사직공원 방향 버스 이용
- 자가용: 서울 종로구 새문안로5가길 32 생산성본부 (주차 유료, 대중교통 권장)

답변 시 유의사항:
- 친근하고 이해하기 쉬운 언어 사용
- 구체적인 정보를 제공하되 불확실한 내용은 담당자(02-3489-475) 문의 안내
- 답변은 한국어로 간결하고 명확하게 작성
- 이모지를 적절히 활용하여 친근감 표현
- 교육 내용에 없는 정보를 추측하여 안내하지 말 것"""


class ChatMessage(BaseModel):
    role: str
    content: str


class ChatRequest(BaseModel):
    messages: list[ChatMessage]


@app.post("/api/chat")
async def chat(req: ChatRequest):
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        raise HTTPException(status_code=503, detail="ANTHROPIC_API_KEY가 설정되지 않았습니다.")

    try:
        import anthropic
        client = anthropic.Anthropic(api_key=api_key)

        messages = [{"role": m.role, "content": m.content} for m in req.messages]

        loop = asyncio.get_event_loop()

        def call_api():
            response = client.messages.create(
                model="claude-haiku-4-5-20251001",
                max_tokens=512,
                system=CURRICULUM_SYSTEM_PROMPT,
                messages=messages,
            )
            return response.content[0].text

        text = await loop.run_in_executor(None, call_api)
        return {"reply": text}

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"챗봇 오류: {str(e)}")


_ensure_credentials_file()


if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", os.environ.get("CHATBOT_PORT", 8001)))
    uvicorn.run("chatbot_server:app", host="0.0.0.0", port=port, reload=False)
